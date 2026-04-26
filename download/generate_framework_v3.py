#!/usr/bin/env python3
"""
Generate 'Beyond the Plan: 3 Pillars, 1 Practice' framework document
Version 3 — Expanded evidence base from 64 classified MRWA Banner Alerts
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ─── Page Setup ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ─── Style Definitions ────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, (size, color, bold) in enumerate([
    (Pt(28), RGBColor(0x1a, 0x3c, 0x5e), True),
    (Pt(20), RGBColor(0x1a, 0x3c, 0x5e), True),
    (Pt(14), RGBColor(0x2c, 0x5f, 0x8a), True),
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.size = size
    h.font.color.rgb = color
    h.font.bold = bold
    h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    h.paragraph_format.space_after = Pt(8)


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('\u2500' * 72)
    run.font.color.rgb = RGBColor(0xbb, 0xbb, 0xbb)
    run.font.size = Pt(8)


def add_quote(doc, text, attribution=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f'\u201c{text}\u201d')
    run.font.italic = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)
    if attribution:
        run2 = p.add_run(f'\n\u2014 {attribution}')
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_styled_table(doc, headers, rows, header_color='1a3c5e', alt_color='f0f4f8'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            if row_idx % 2 == 1:
                set_cell_shading(cell, alt_color)
    return table


def add_callout(doc, text, label=None):
    """Add a highlighted callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    if label:
        run = p.add_run(f'{label}: ')
        run.bold = True
        run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x5e)
        run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BEYOND THE PLAN')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x5e)
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('3 Pillars, 1 Practice')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)
run.font.name = 'Calibri'

add_horizontal_rule(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A Frontline Safety Framework\nfor Temporary Traffic Management')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Industry Position Paper')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'April 2026 \u2022 Version 3 \u2022 Draft for Review')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared for industry consultation\nwith Main Roads Western Australia and traffic management organisations')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.italic = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'Temporary traffic management in Western Australia is supported by comprehensive standards \u2014 '
    'the AGTTM, MRWA Code of Practice, and AS 1742.3 provide detailed guidance for planning, designing, '
    'and implementing traffic management at work zones. These standards represent significant expertise '
    'and are fit for their purpose: describing how a job should be set up and how traffic should be managed.'
)

doc.add_paragraph(
    'Yet workplace incidents in traffic management continue to occur on jobs where the paperwork was '
    'complete, the crew was accredited, and the signs were on the ute. The gap is not in the standards. '
    'The gap is between what the plan describes at the start of a shift and what the crew experiences '
    'during it. In safety science, this is known as the gap between "work as imagined" and "work as done" '
    '\u2014 and it is a gap that compliance-based systems cannot close.'
)

doc.add_paragraph(
    'This paper proposes a frontline safety framework \u2014 Beyond the Plan: 3 Pillars, 1 Practice \u2014 '
    'that names and structures what experienced traffic controllers already know: that the three things '
    'you verify before you start (competent crew, sound plan, right equipment) are necessary but not '
    'sufficient. The fourth element \u2014 dynamic awareness, the continuous practice of reading conditions '
    'and being willing to act on what you see \u2014 is what separates a safe crew from a crew that had '
    'all the right paperwork but still got hurt.'
)

doc.add_paragraph(
    'The framework draws on two complementary traditions in safety science. The three pillars are grounded '
    'in Safety I thinking \u2014 the established approach of identifying hazards, implementing controls, '
    'and verifying compliance. Dynamic awareness is grounded in Safety II thinking \u2014 the emerging '
    'paradigm that focuses on building adaptive capacity, studying how work succeeds, and treating workers '
    'as the solution rather than the variable to control. The framework does not choose between these '
    'traditions. It uses each where it is strongest: Safety I for the preconditions, Safety II for the '
    'practice that sustains them.'
)

doc.add_paragraph(
    'The framework is grounded in evidence from MRWA Banner Alerts, mapped against existing legislative '
    'and regulatory requirements, and designed for practical adoption without adding compliance burden. '
    'It does not replace current standards. It provides the connecting tissue that makes them work in '
    'real time, on the road, when conditions change.'
)

p = doc.add_paragraph()
run = p.add_run('Key proposition: ')
run.bold = True
p.add_run(
    'The three pillars are preconditions \u2014 without them, the job should not proceed. '
    'The practice is continuous \u2014 without it, the preconditions degrade as conditions change. '
    'Communication is the thread that holds all four together. And the organisation is accountable '
    'for creating the conditions where all of this is possible.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 2. THE PROBLEM
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('2. The Problem: When the Paperwork Is Perfect', level=1)

doc.add_heading('2.1 The Gap Between Plan and Reality', level=2)

doc.add_paragraph(
    'A Traffic Management Plan is a snapshot. It describes the work zone as it was designed \u2014 the sign '
    'positions, the speed zones, the traffic control points. It assumes the weather stays the same, the '
    'traffic behaves predictably, and the hazards are known in advance. These assumptions hold at the '
    'moment the plan is written. They do not necessarily hold at 2:30 PM on a Tuesday when the wind '
    'picks up, the sun shifts, and the traffic volume doubles because a nearby school just finished for the day.'
)

doc.add_paragraph(
    'This is not a failure of the planning process. Plans cannot predict every condition. The purpose of '
    'a plan is to establish a safe baseline. The problem arises when the baseline is treated as the '
    'complete picture \u2014 when the crew follows the plan right up to the moment someone gets hurt, '
    'and the investigation reveals that conditions had changed and nobody adjusted.'
)

doc.add_paragraph(
    'The safety science literature has a name for this gap. Erik Hollnagel, one of the founders of '
    'Resilience Engineering, distinguishes between "work as imagined" \u2014 what the procedure, the plan, '
    'and the risk assessment assume will happen \u2014 and "work as done" \u2014 what actually happens when '
    'real people do real work in real conditions. The gap between the two is not an anomaly. It is a '
    'permanent feature of complex work. The question is not whether the gap exists, but whether the '
    'system is designed to cope with it.'
)

doc.add_heading('2.2 The Evidence: Banner Alert Case Studies', level=2)

doc.add_paragraph(
    'The MRWA Banner Alert system provides documented evidence of incidents that occurred on jobs '
    'where the standard preconditions were met. The following cases illustrate the gap:'
)

# Case study 1
doc.add_heading('Case 1: TC Slips on Soft Roadside Batter (EQ#57872)', level=3)
doc.add_paragraph(
    'A Traffic Controller exited their vehicle on the passenger side at 6:15 AM. Their feet slipped '
    'on a soft batter at the road shoulder. They grabbed the door and roll bar to right themselves, '
    'sustaining a shoulder injury that resulted in a Lost Time Injury.'
)
p = doc.add_paragraph()
run = p.add_run('Pre-start status: ')
run.bold = True
p.add_run(
    'The crew was accredited. The TMP was in place. The signs were loaded and positioned. '
    'All three pillars were satisfied.'
)
p = doc.add_paragraph()
run = p.add_run('What the plan didn\'t account for: ')
run.bold = True
p.add_run(
    'Low light at dawn combined with a soft roadside batter that looked stable but wasn\'t. '
    'This hazard only existed in that moment \u2014 it was not present when the site was assessed '
    'in daylight, and it would not be present after the ground dried. Dynamic awareness \u2014 '
    'scanning the ground surface before stepping out, especially in low light \u2014 would have '
    'identified the hazard before it caused injury.'
)

# Case study 2
doc.add_heading('Case 2: Worker Trips on Faded Wheel Stop (EQ#57539)', level=3)
doc.add_paragraph(
    'A worker tripped over a parking bay wheel stop bar that had faded over time, losing its '
    'visual contrast with the surrounding surface. The worker was not aware of the tripping hazard '
    'and sustained a leg injury.'
)
p = doc.add_paragraph()
run = p.add_run('Pre-start status: ')
run.bold = True
p.add_run('The wheel stop was there the entire time. The site had been in use for months.')
p = doc.add_paragraph()
run = p.add_run('What the plan didn\'t account for: ')
run.bold = True
p.add_run(
    'The gradual degradation of a safety feature into a hazard. The wheel stop\'s colour markings '
    'had faded, reducing its visibility to the point where it became a trip hazard rather than a '
    'safety device. Nobody noticed because it blended in. Dynamic awareness \u2014 scanning for '
    'trip hazards on arrival and flagging them \u2014 would have caught this before someone was hurt.'
)

# Case study 3
doc.add_heading('Case 3: Sun Glare Changes Mid-Job (EQ#57762)', level=3)
doc.add_paragraph(
    'A traffic management job started with clear visibility. As the afternoon progressed, the sun '
    'shifted to a position that caused direct glare into the Traffic Controller\'s line of sight. '
    'The TC could no longer see oncoming traffic clearly, creating a significant safety risk.'
)
p = doc.add_paragraph()
run = p.add_run('Pre-start status: ')
run.bold = True
p.add_run('The TMP was correct for the conditions at the start of the shift. Signs were positioned properly.')
p = doc.add_paragraph()
run = p.add_run('What the plan didn\'t account for: ')
run.bold = True
p.add_run(
    'The sun\'s position at 3:30 PM. The plan was written for the setup time, not the entire shift. '
    'Dynamic awareness \u2014 recognising that visibility had degraded and triggering a reassessment \u2014 '
    'would have led to repositioning the TC, adding a spotter, or pausing until the glare passed.'
)

# Case study 4
doc.add_heading('Case 4: Concrete Saw Strikes Worker (EQ#56820)', level=3)
doc.add_paragraph(
    'A quick-cut saw struck a worker in the face while cutting reinforced concrete pipe. The SWMS '
    'did not contain information on correct placement of the pipe to prevent pinching the blade. '
    'The prestart meeting form had been reused across multiple days instead of being completed fresh each day.'
)
p = doc.add_paragraph()
run = p.add_run('Pre-start status: ')
run.bold = True
p.add_run(
    'A SWMS existed. A prestart form was completed. The crew had qualifications. On paper, the '
    'preconditions were met.'
)
p = doc.add_paragraph()
run = p.add_run('What the plan didn\'t account for: ')
run.bold = True
p.add_run(
    'The SWMS was incomplete for the specific task. The prestart form was a formality rather than a '
    'genuine hazard assessment. This is a failure at both the organisational level (the system allowed '
    'a generic SWMS and reused prestart forms) and the crew level (nobody questioned whether the SWMS '
    'covered the work they were about to do). Dynamic awareness would have prompted the question: '
    '"Does our plan actually cover what we\'re doing right now?"'
)

doc.add_heading('2.3 The Pattern', level=2)

doc.add_paragraph(
    'Across these incidents, a consistent pattern emerges:'
)

add_bullet(doc, 'The three preconditions were satisfied on paper', 'Compliance: ')
add_bullet(doc, 'Conditions changed during the job, or a hazard existed that was not identified in the plan', 'Change: ')
add_bullet(doc, 'The crew did not adjust \u2014 either because they didn\'t notice the change, didn\'t feel authorised to act, or didn\'t have a process for reassessment', 'Gap: ')

doc.add_paragraph(
    'This pattern is not unique to Western Australia. It is consistent with safety research across '
    'high-risk industries, from construction to aviation to offshore operations. The pattern has a name '
    'in the safety literature: the gap between "work as imagined" (what the plan assumes) and "work as '
    'done" (what actually happens). The question is not whether this gap exists \u2014 it always does. '
    'The question is whether the crew has the awareness, the authority, and the process to bridge it.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. THE INTELLECTUAL FOUNDATION
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('3. The Intellectual Foundation: From Safety I to Safety II', level=1)

doc.add_paragraph(
    'This framework does not emerge from nowhere. It sits at the intersection of two traditions in '
    'safety science \u2014 one established, one emerging. Understanding both is essential because the '
    '3 Pillars draw on the first tradition, while the Practice draws on the second. The framework\'s '
    'strength is in using each tradition where it is strongest.'
)

doc.add_heading('3.1 Safety I: The Established Tradition', level=2)

doc.add_paragraph(
    'Safety I is the approach most safety professionals were trained in. It underpins the DuPont safety '
    'system, the MRWA Code of Practice, and virtually every compliance framework in Australian traffic '
    'management. Its core assumption is that safety is the absence of negative outcomes \u2014 no injuries, '
    'no incidents, no non-conformances. When something goes wrong, Safety I asks: what broke? What caused '
    'the failure? How do we prevent it from happening again?'
)

doc.add_paragraph(
    'The DuPont safety system, developed over more than 200 years since the company began manufacturing '
    'gunpowder in 1802, is the most influential expression of Safety I thinking in industrial safety. '
    'DuPont\'s Front Line Management course, widely adopted across high-risk industries, teaches '
    'supervisors to manage safety through behavioural observation, compliance auditing, and root-cause '
    'investigation. The Bradley Curve \u2014 DuPont\'s cultural maturity model describing the progression '
    'from Reactive through Dependent and Independent to Interdependent \u2014 remains the most widely '
    'recognised framework for assessing safety culture maturity.'
)

doc.add_paragraph(
    'Safety I has achieved remarkable results. Industries that adopted it \u2014 chemical manufacturing, '
    'oil and gas, aviation \u2014 saw dramatic reductions in workplace fatalities and serious injuries. '
    'In traffic management, Safety I gives us the prestart checklist, the TMP, the SWMS, and the audit '
    'regime. These are not trivial achievements. They represent a baseline of systematic safety practice '
    'that prevents many categories of harm.'
)

doc.add_heading('3.2 The Limitations of Safety I in Traffic Management', level=2)

doc.add_paragraph(
    'Safety I was developed for controlled environments \u2014 chemical plants, factories, refineries '
    '\u2014 where the workplace is enclosed, the workforce is permanent, and conditions are stable. '
    'Temporary traffic management is fundamentally different in four ways:'
)

add_bullet(doc, 'The environment changes constantly. DuPont plants are the same every day. A TC work zone is different every setup. Weather, traffic, light, road surface, and adjacent construction activity can all change within a single shift.', 'Uncontrolled environment: ')
add_bullet(doc, 'The public drives through the workplace at 80 km/h. DuPont controls who enters the plant and how they behave. TCs have no control over the motorists who pass through their work zone, often at speed and sometimes under the influence, distracted, or confused by the layout.', 'Public in the hazard zone: ')
add_bullet(doc, 'TCs set up, work, and tear down, sometimes multiple times a day. There is no "steady state" to optimise. Each setup is a new configuration with new risks.', 'Mobile and temporary work: ')
add_bullet(doc, 'The workforce is often labour-hire, short-term, and high-turnover. DuPont has career employees who develop deep institutional knowledge over decades. Many TCs are on their first or second season, working alongside people they met this morning.', 'Casualised workforce: ')

doc.add_paragraph(
    'These differences matter because they mean that the gap between "work as imagined" and "work as done" '
    'is wider, more variable, and less predictable in traffic management than in the environments Safety I '
    'was designed for. A compliance system that works in a chemical plant \u2014 where conditions are '
    'stable and deviations are rare \u2014 struggles in a traffic management context where conditions '
    'change continuously and deviation from the plan is often the only way to stay safe.'
)

doc.add_heading('3.3 Safety II: The Emerging Paradigm', level=2)

doc.add_paragraph(
    'Since approximately 2014, a new paradigm has emerged in safety science that directly addresses the '
    'limitations of Safety I. Three thinkers and their frameworks form the core of this shift:'
)

doc.add_heading('Erik Hollnagel: Safety II and Resilience Engineering', level=3)

doc.add_paragraph(
    'Hollnagel\'s foundational question is the simplest and most profound: instead of asking "what went '
    'wrong?" why don\'t we ask "what goes right?" His insight is that safety is not the absence of '
    'negatives \u2014 it is the presence of capacities. A crew that successfully navigates a shift '
    'where the wind blows signs over, traffic ignores speed zones, and the sun blinds the TC is '
    'demonstrating safety. Not because nothing bad happened, but because they had the capacity to '
    'adapt and succeed.'
)

headers = ['', 'Safety I (Traditional)', 'Safety II (Hollnagel)']
rows = [
    ['Core question', 'What went wrong?', 'What goes right?'],
    ['Safety is defined as', 'Absence of negatives (no injuries)', 'Presence of adaptive capacity'],
    ['Humans are', 'The variable to control', 'The reason things go right'],
    ['Approach', 'Prevent things from going wrong', 'Ensure things go right more often'],
    ['After an incident', 'Find what broke', 'Understand how people were trying to succeed'],
]
add_styled_table(doc, headers, rows)

doc.add_paragraph('')
doc.add_paragraph(
    'Hollnagel\'s Resilience Engineering framework identifies four capacities that enable systems to '
    'cope with surprise: the ability to respond to what happens, to monitor what is happening, to '
    'anticipate what may happen, and to learn from what has happened. These four capacities map '
    'directly to the components of Dynamic Awareness: reading conditions (monitor), adjusting controls '
    '(respond), experience-based expectation (anticipate), and post-job review (learn).'
)

doc.add_heading('Sidney Dekker: Safety Differently', level=3)

doc.add_paragraph(
    'Dekker builds on Hollnagel with three provocations that challenge the foundations of compliance-based '
    'safety. First, safety is not defined by the absence of negatives \u2014 it is defined by the presence '
    'of capacities. Second, safety is not a bureaucratic accountability \u2014 it is an ethical '
    'responsibility. Third, and most importantly, people are not the problem to control \u2014 they are '
    'the resource to harness.'
)

doc.add_paragraph(
    'Dekker\'s most significant practical contribution is the concept of Restorative Just Culture. When '
    'something goes wrong, compliance-based systems ask "who violated the rule?" Dekker asks "who has been '
    'hurt, and what do they need?" The shift from justice as retribution to justice as restoration is '
    'profound. In traffic management, it means moving from a culture where the crew leader who calls a stop '
    'gets asked "why aren\'t you following the plan?" to a culture where they get asked "what did you see '
    'that the plan didn\'t account for?"'
)

doc.add_paragraph(
    'Dekker\'s 2018 book "The Safety Anarchist" argues that more rules, more paperwork, and more '
    'bureaucracy can actually reduce safety by pushing risk underground and removing workers\' ability to '
    'adapt. This is directly observable in TC operations: when the prestart becomes a form to complete '
    'rather than a conversation to have, it stops serving its safety purpose and becomes a compliance '
    'exercise that teaches crews that the paperwork matters more than the practice.'
)

doc.add_heading('Todd Conklin: Human and Organisational Performance (HOP)', level=3)

doc.add_paragraph(
    'Conklin is the most practical of the three. His Five Principles of Human Performance provide a '
    'direct operational framework:'
)

add_bullet(doc, 'Humans make mistakes, even well-trained, well-meaning humans. Stop designing systems that assume they won\'t.', 'Error is normal: ')
add_bullet(doc, 'Punishing the person doesn\'t fix the system that allowed the error.', 'Blame fixes nothing: ')
add_bullet(doc, 'People do what the system incentivises, not what the procedure says.', 'Systems drive behaviour: ')
add_bullet(doc, 'Don\'t just investigate failures. Study how work actually happens day-to-day.', 'Learn from normal work: ')
add_bullet(doc, 'Accidents come from normal variations that combine differently this time, not from broken parts.', 'Failure is resurfacing normal variability: ')

doc.add_paragraph(
    'Conklin\'s key practical tools are Learning Teams and Pre-Accident Investigations. Learning Teams '
    'replace root-cause investigations with facilitated conversations about how work actually gets done. '
    'Pre-Accident Investigations study normal work before failures occur. Both are directly applicable '
    'to TC crews: instead of waiting for an incident to investigate, sit down with the crew after a '
    'normal shift and ask "what did you have to do today to make this work?"'
)

doc.add_heading('3.4 The Structural Contradiction in Australian Traffic Management', level=2)

doc.add_paragraph(
    'There is a contradiction at the heart of Australian traffic management safety that the framework '
    'is designed to address. Australia is a global leader in Safe System thinking for road users. The '
    'National Road Safety Strategy 2021-2030, Austroads frameworks, and MRWA\'s own road design '
    'standards all embrace the principle that humans are fallible and systems should be designed to '
    'tolerate error. This is Safety II thinking applied to the public.'
)

doc.add_paragraph(
    'But for the workers who set up and manage traffic control? The approach remains overwhelmingly '
    'Safety I: prescriptive TMPs, compliance audits, behavioural observation, root-cause investigation. '
    'No Australian traffic management company or road authority appears to have formally adopted Safety '
    'Differently, Safety II, or HOP as their primary safety framework for traffic workers.'
)

doc.add_paragraph(
    'This creates a structural contradiction: the road safety strategy says "humans make mistakes, '
    'design the system to cope," but the safety management of the people implementing that strategy '
    'says "comply with the plan, don\'t deviate, all injuries are preventable." The 3 Pillars, 1 '
    'Practice framework bridges this contradiction. The three pillars provide the Safety I baseline '
    '\u2014 verifiable, auditable, necessary. Dynamic awareness provides the Safety II practice \u2014 '
    'continuous, adaptive, behavioural. Together, they apply the same Safe System principle to workers '
    'that Australia already applies to road users.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. THE FRAMEWORK
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('4. The Framework: 3 Pillars, 1 Practice', level=1)

doc.add_paragraph(
    'The framework is built on a single observation: the things you verify before a job starts are '
    'different in kind from the thing you do while the job is happening. The three pillars are static \u2014 '
    'they are preconditions that, once verified, remain stable for the duration of the task. They are '
    'grounded in Safety I thinking \u2014 they are auditable, verifiable, and compliance-oriented. The '
    'practice is continuous \u2014 it begins when you arrive and does not stop until you pack down. It '
    'is grounded in Safety II thinking \u2014 it is behavioural, adaptive, and capacity-oriented.'
)

add_horizontal_rule(doc)

# Pillar 1
doc.add_heading('Pillar 1: Competent Crew', level=2)

p = doc.add_paragraph()
run = p.add_run('The question: ')
run.bold = True
run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)
p.add_run('Do we have people who know what they\'re doing?')

doc.add_paragraph(
    'Competence is not the same as accreditation. A Traffic Controller with a valid BWTM card has met '
    'the minimum training requirement, but competence on a specific job requires more: familiarity with '
    'the road, experience with the type of setup, and the judgement that comes from having worked in '
    'conditions similar to those expected on the day.'
)

doc.add_paragraph(
    'The distinction matters because accreditation is binary (you have the card or you don\'t) but '
    'competence is contextual. A TC who is highly competent on metropolitan signalised intersections '
    'may not be competent on a rural highway at night. A TC who has installed hundreds of lane closures '
    'may not be competent on a mobile convoy operation. Competence is always relative to the task, '
    'the location, and the conditions.'
)

p = doc.add_paragraph()
run = p.add_run('What this pillar requires:')
run.bold = True

add_bullet(doc, 'All crew members hold current accreditations appropriate to their role')
add_bullet(doc, 'At least one crew member has experience with the specific type of setup required')
add_bullet(doc, 'The crew has been briefed on the site-specific risks, not just the generic hazards')
add_bullet(doc, 'Any crew member new to this type of work is paired with an experienced hand')

p = doc.add_paragraph()
run = p.add_run('It\'s broken when: ')
run.bold = True
run.font.color.rgb = RGBColor(0xcc, 0x33, 0x33)
p.add_run(
    'The crew is accredited but nobody has worked this road before, or a TC is assigned to a setup '
    'type they\'ve never done without supervision.'
)

# Pillar 2
doc.add_heading('Pillar 2: Sound Plan', level=2)

p = doc.add_paragraph()
run = p.add_run('The question: ')
run.bold = True
run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)
p.add_run('Do we know what we\'re doing and how?')

doc.add_paragraph(
    'A sound plan is one that is specific to the site, current for the conditions, and understood by '
    'the crew. It is not a generic TGS pulled from a folder. It is not a TMP that was written for a '
    'different time of day, a different traffic volume, or a different road configuration. And it is '
    'not a document that sits on the dashboard unread while the crew sets up from memory.'
)

doc.add_paragraph(
    'The MRWA Code of Practice and AGTTM provide the design standards for traffic management plans. '
    'This pillar is not about the quality of those standards \u2014 it is about whether the plan that '
    'reaches the crew is fit for the specific job they are about to do. A plan that was correct when '
    'it was written but does not account for current conditions is not a sound plan for today.'
)

p = doc.add_paragraph()
run = p.add_run('What this pillar requires:')
run.bold = True

add_bullet(doc, 'The TMP and TGS are site-specific, not generic')
add_bullet(doc, 'The plan accounts for the time of day, expected traffic volumes, and known environmental conditions')
add_bullet(doc, 'The crew has read and understood the plan, not just received it')
add_bullet(doc, 'The plan includes provisions for conditions that may change during the shift (weather, light, traffic)')

p = doc.add_paragraph()
run = p.add_run('It\'s broken when: ')
run.bold = True
run.font.color.rgb = RGBColor(0xcc, 0x33, 0x33)
p.add_run(
    'The TMP is generic and doesn\'t match site conditions, the TGS was written for a different setup, '
    'or the crew hasn\'t actually read the plan.'
)

# Pillar 3
doc.add_heading('Pillar 3: Right Equipment', level=2)

p = doc.add_paragraph()
run = p.add_run('The question: ')
run.bold = True
run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)
p.add_run('Do we have what the plan requires?')

doc.add_paragraph(
    'Having the right equipment means more than loading the signs listed on the TGS schedule. It means '
    'the signs are serviceable (not faded, damaged, or wrong-sized), the supplementary equipment is '
    'available (speed feedback signs, lighting towers, night wands), and there are sufficient reserves '
    'to adapt if conditions change during the shift.'
)

doc.add_paragraph(
    'The MRWA Code of Practice Section 6.16 mandates speed feedback signs at static worksites on Main '
    'Roads where workers on foot are not protected by barriers. Section 6.14 requires that all signs '
    'and devices are maintained in a serviceable condition throughout the work period. This pillar is '
    'about verifying compliance with these requirements before the job starts, not discovering their '
    'absence during it.'
)

p = doc.add_paragraph()
run = p.add_run('What this pillar requires:')
run.bold = True

add_bullet(doc, 'All signs and devices required by the TGS are loaded and in serviceable condition')
add_bullet(doc, 'Supplementary equipment specified in the TMP is available (speed feedback signs, lighting, communication equipment)')
add_bullet(doc, 'Equipment for contingencies is available (spare signs, additional delineation, wet weather gear)')
add_bullet(doc, 'The crew has verified the equipment list against the plan, not just loaded from habit')

p = doc.add_paragraph()
run = p.add_run('It\'s broken when: ')
run.bold = True
run.font.color.rgb = RGBColor(0xcc, 0x33, 0x33)
p.add_run(
    'Signs are loaded but the wrong size or faded beyond readability, the speed feedback sign that '
    'the TMP requires is back at the depot, or the crew loaded from memory without checking the schedule.'
)

add_horizontal_rule(doc)

# The Practice
doc.add_heading('The Practice: Dynamic Awareness', level=2)

add_quote(doc,
    'Dynamic Awareness is the practice of continuously reading the job while it\'s happening \u2014 '
    'and being willing to act on what you see.',
    'Beyond the Plan Framework'
)

doc.add_paragraph(
    'Dynamic awareness is not a checklist. It is not something you tick off at the pre-start brief. '
    'It is the only element of the framework that is active for the entire duration of the job, from '
    'the moment the first crew member steps out of the vehicle to the moment the last sign is loaded '
    'back onto the ute.'
)

doc.add_paragraph(
    'The three pillars are static \u2014 once verified, they remain stable. Dynamic awareness is '
    'continuous. It exists precisely because the three pillars, no matter how well they are satisfied '
    'at the start, will degrade as conditions change. The wind picks up. The sun shifts. The traffic '
    'volume changes. A crew member fatigues. A hazard that wasn\'t there yesterday appears. The plan '
    'doesn\'t cover it. That\'s where dynamic awareness becomes the safety net.'
)

doc.add_paragraph(
    'Dynamic awareness is a Safety II concept. Where Safety I asks "did you follow the procedure?" '
    'Dynamic Awareness asks "what did you see, and what did you do about it?" Where Safety I treats '
    'deviation from the plan as non-compliance, Dynamic Awareness treats it as adaptation \u2014 '
    'provided it is informed, proportionate, and communicated. The MRWA TMP risk register section '
    '3.2.1.2 already authorises this: "Traffic Management Personnel may adjust or change the site '
    'where a significant safety hazard is identified." The framework gives this buried sentence '
    'prominence and process.'
)

doc.add_heading('The Three Components of Dynamic Awareness', level=3)

p = doc.add_paragraph()
run = p.add_run('1. Reading Conditions')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)

doc.add_paragraph(
    'This is the observational foundation. Reading conditions means paying attention to what is actually '
    'happening, not what the forecast said this morning or what the TMP assumed. It includes monitoring '
    'weather changes (wind, rain, visibility), traffic behaviour (are drivers slowing down? is there '
    'congestion building?), road surface conditions (wet, soft, dusty), light conditions (sun glare, '
    'dusk, artificial lighting adequacy), and crew state (fatigue, distraction, morale).'
)

doc.add_paragraph(
    'Reading conditions is not passive observation. It is active scanning \u2014 deliberately looking '
    'for what has changed, not just what is there. The experienced TC who notices that the wind has '
    'strengthened enough to make a large sign unstable is reading conditions. The TC who walks the '
    'site on arrival and spots a pothole that wasn\'t there last week is reading conditions. The '
    'supervisor who sees a crew member struggling with fatigue at the two-hour mark is reading conditions.'
)

doc.add_paragraph(
    'In Hollnagel\'s Resilience Engineering framework, reading conditions corresponds to the "monitor" '
    'capacity \u2014 the ability to know what to look for and to recognise when something is changing. '
    'It is the foundation of adaptive capacity. Without it, adjustment is impossible because you don\'t '
    'know what you\'re adjusting to.'
)

p = doc.add_paragraph()
run = p.add_run('2. Adjusting Controls')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)

doc.add_paragraph(
    'When conditions change, do you actually change what you\'re doing? Or do you push through because '
    '"the TMP says so"? Adjusting controls means having both the authority and the willingness to '
    'modify the traffic management setup in response to what you observe.'
)

doc.add_paragraph(
    'This is where the framework connects directly to existing standards. The MRWA TMP risk register '
    'section 3.2.1.2 states: "Traffic Management Personnel may adjust or change the site where a '
    'significant safety hazard is identified with appropriate approval and risk assessment." This '
    'sentence, buried in a risk register, is the regulatory basis for dynamic awareness. The framework '
    'gives it prominence and process.'
)

doc.add_paragraph(
    'Adjusting controls might mean repositioning a TC who can\'t see oncoming traffic due to sun glare. '
    'It might mean reducing the speed zone when drivers are ignoring the posted limit. It might mean '
    'adding a spotter when visibility drops. It might mean stopping the job entirely until conditions '
    'improve. The adjustment is always proportional to the observed change. In Hollnagel\'s framework, '
    'this corresponds to the "respond" capacity \u2014 knowing what to do when something happens.'
)

p = doc.add_paragraph()
run = p.add_run('3. Stop Authority')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)

doc.add_paragraph(
    'The willingness to call a halt. Not after something goes wrong \u2014 when you feel something is '
    'off. This is where culture matters more than paperwork. If a TC feels they cannot call a stop '
    'without being reprimanded, the entire practice collapses. Stop authority is not a power \u2014 it '
    'is a permission that must be granted by the organisation, reinforced by supervisors, and '
    'exercised by every crew member without consequence.'
)

doc.add_paragraph(
    'The DuPont safety culture model operates on a principle that any worker can stop any job at any '
    'time without question. This is not a sign of weakness or disruption \u2014 it is a sign of a '
    'mature safety culture. In Conklin\'s HOP framework, stop authority is an expression of the '
    'principle that error is normal and systems must be designed to tolerate it. When a TC calls a '
    'stop, they are not causing disruption \u2014 they are exercising the adaptive capacity that '
    'prevents a near-miss from becoming an incident.'
)

add_horizontal_rule(doc)

# The Thread
doc.add_heading('The Thread: Communication', level=2)

doc.add_paragraph(
    'Communication is not a fourth pillar or a fifth element. It is the thread that runs through all '
    'four components of the framework. The three pillars are verified through communication: pre-start '
    'briefings, equipment checks, plan reviews. Dynamic awareness is practised through communication: '
    'calling out what you see, radioing changes, speaking up about concerns, and \u2014 critically \u2014 '
    'calling a stop when something feels wrong.'
)

doc.add_paragraph(
    'Communication failures appear in banner alerts not as standalone causes but as contributing factors '
    'that amplify other failures. In EQ#56820, the prestart form was a communication failure \u2014 a '
    'form completed without genuine engagement. In EQ#57801, the chain recovery was a communication '
    'failure \u2014 nobody flagged that proper recovery equipment was not available. In each case, '
    'communication was the mechanism that could have bridged the gap between what was planned and what '
    'was actually happening.'
)

doc.add_paragraph(
    'For this reason, communication is not positioned as a separate pillar. Making it one would risk '
    'reducing it to a checkbox \u2014 "radio check, tick" \u2014 rather than recognising it as the '
    'continuous process that enables every other element of the framework. The DuPont FLM course '
    'teaches the "5-Minute Conversation" \u2014 short, focused safety conversations that are not '
    'checklists but genuine observations: "I noticed you repositioned the taper \u2014 what did you '
    'see?" This is communication as a skill to teach and practise, not a box to tick.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 5. THE CULTURAL REPRODUCTION PROBLEM
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('5. The Cultural Reproduction Problem', level=1)

doc.add_paragraph(
    'The hardest question facing any safety framework is not "what should good practice look like?" '
    'but "how do you change practice when the people transmitting the culture are the problem?" This '
    'section addresses that question directly, because no framework \u2014 Safety I, Safety II, or '
    'anything else \u2014 can succeed if the people leading crews are modelling the wrong behaviour.'
)

doc.add_heading('5.1 The Apprenticeship of Bad Practice', level=2)

doc.add_paragraph(
    'When a new Traffic Controller arrives on their first job, they are a blank slate. They want to do '
    'a good job, but they don\'t yet know what good looks like. What they learn is whatever they see. '
    'If the crew leader does the prestart properly, walks the site, discusses hazards, and monitors '
    'conditions throughout the shift, the new TC learns that this is how professionals work. If the '
    'crew leader mumbles through the prestart, takes a phone call halfway through, doesn\'t walk the '
    'site, and spends the shift on their phone, the new TC learns something very different: this is '
    'how professionals work.'
)

doc.add_paragraph(
    'The problem is not that the new TC is lazy or disengaged. The problem is that they are doing '
    'exactly what humans do: learning from observation. They assume the person leading them is the '
    'standard. And why wouldn\'t they? The crew leader has more experience, more tickets, and more '
    'authority. If this is how they do it, this must be how it\'s done.'
)

doc.add_paragraph(
    'Over a year, a single disengaged crew leader might influence 20-30 different TCs who rotate '
    'through their crew. Each of those takes the model to their next crew. Within three years, one '
    'crew leader\'s minimum-standard approach has infected an entire labour pool. This is not a training '
    'problem \u2014 it is a cultural reproduction problem, and it is exponential.'
)

doc.add_heading('5.2 The Normalisation of Deviance', level=2)

doc.add_paragraph(
    'The sociologist Diane Vaughan coined the term "normalisation of deviance" in her study of the '
    'Challenger space shuttle disaster. She did not mean people were deliberately deviating from '
    'standards. She meant that gradually, incrementally, what was once recognised as substandard '
    'becomes accepted as normal, because everyone sees everyone else doing it and nobody says anything. '
    'The standard doesn\'t drop in one step. It erodes.'
)

doc.add_paragraph(
    'In TC work, the erosion looks like this: the original standard is a full prestart with a site '
    'walk, hazard discussion, and continuous monitoring. The first erosion shortens the prestart and '
    'replaces the site walk with a drive-past. The second erosion makes the prestart a formality '
    'and drops the hazard discussion. The third erosion turns the prestart into a checkbox and '
    'replaces monitoring with "text me if something happens." The new TC who arrives at the third '
    'erosion thinks this is the standard. They have never seen the original.'
)

doc.add_heading('5.3 Why Compliance Systems Make It Worse', level=2)

doc.add_paragraph(
    'Here is the most perverse aspect of the current system: compliance mechanisms actually accelerate '
    'the normalisation of deviance. When a crew leader ticks the prestart box without genuinely '
    'conducting the prestart, two things happen simultaneously. First, the record says it was done '
    '\u2014 the audit trail is clean, the company sees compliance, no flags are raised. Second, the '
    'TC sees it being faked \u2014 and they learn that the paperwork matters more than the practice. '
    'They learn that safety is performative. They learn that the system can be gamed, and that '
    'everyone knows it, and that nobody cares as long as the form is signed.'
)

doc.add_paragraph(
    'This is worse than having no system at all. A system that can be gamed teaches people to game '
    'systems. It trains cynicism. And cynical TCs do not practise Dynamic Awareness because they '
    'have already concluded that safety is theatre. More checklists, more audits, and more forms '
    'will not fix this \u2014 the disengaged crew leader will tick those boxes too. Compliance is '
    'the language of minimums, and minimums are what got us here.'
)

doc.add_heading('5.4 What the Framework Does Differently', level=2)

doc.add_paragraph(
    'The 3 Pillars, 1 Practice framework addresses cultural reproduction in three ways that the '
    'current compliance system does not:'
)

add_bullet(doc, 'The three pillars are verification questions, not checkbox declarations. "Can you demonstrate the skills for this specific site?" cannot be faked the way "is the crew competent?" can be ticked. Verification requires demonstration, not declaration.', 'The 3 Pillars create a visible floor: ')
add_bullet(doc, 'Dynamic Awareness gives the new TC something to look up to. Right now, the new TC has no model of what good looks like beyond the checklist. Dynamic Awareness says: "The standard isn\'t \'set up and stand there.\' The standard is \'set up, then keep reading the job.\' Here\'s what that looks like." Even if the person next to them isn\'t hitting that standard, the new TC at least has a conceptual target.', 'Dynamic Awareness creates a visible ceiling: ')
add_bullet(doc, 'The framework explicitly holds organisations accountable for ensuring that crew leaders can model the ceiling, not just meet the floor. If a crew leader cannot demonstrate Dynamic Awareness, the organisation has failed \u2014 not the individual TC who copies them. This is a fundamental shift from the current system, which holds individuals accountable for compliance but never holds the organisation accountable for the conditions that produce non-compliance.', 'The organisational layer creates accountability for modelling: ')

# ═══════════════════════════════════════════════════════════════════════════════
# 6. MONITORING WHAT MATTERS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('6. Monitoring What Matters: From Compliance Audits to Behavioural Visibility', level=1)

doc.add_paragraph(
    'If the framework is going to work, there must be a way to know whether it is being practised. '
    'But here is the central challenge: bad behaviour is easy to monitor, good behaviour is hard. '
    'A cone out of position can be photographed. A sign facing the wrong way can be ticked as a '
    'non-conformance. But how do you monitor vigilance? How do you audit situational awareness? '
    'How do you record the fact that a TC was actively scanning the approach every 30 seconds?'
)

doc.add_heading('6.1 Why Monitoring Good Behaviour Is Different', level=2)

doc.add_paragraph(
    'The DuPont behavioural observation programs that were widely adopted in the 1990s and 2000s '
    'often failed because they turned behavioural safety into a surveillance exercise. Workers '
    'performed for the observer rather than working naturally. The observation changed the behaviour '
    'being observed. This is a well-documented phenomenon in safety research, and it explains why '
    'checklist-based behavioural observation programs produce data but not insight.'
)

doc.add_paragraph(
    'Safety II offers a different approach. Instead of monitoring for compliance, monitor for capacity. '
    'Instead of auditing what went wrong, study what goes right. Instead of watching workers, '
    'listen to them. The practical techniques below are drawn from Hollnagel, Conklin, and Dekker, '
    'adapted for the specific conditions of temporary traffic management.'
)

doc.add_heading('6.2 Practical Techniques for TC Operations', level=2)

doc.add_heading('The "What Did You See?" Conversation', level=3)

doc.add_paragraph(
    'Not an audit. Not a checklist. A genuine question, asked with curiosity, not judgment: "I noticed '
    'you moved the stop/slow downstream \u2014 what did you see?" "You called for a hold before that '
    'truck came through \u2014 what told you to do that?" The answers reveal the Dynamic Awareness '
    'process. The act of asking validates it. Over time, this creates a culture where TCs expect to '
    'be asked about their thinking, not just their compliance. This is not a management function \u2014 '
    'it is a crew function. When peers ask each other, it builds interdependence. When only '
    'supervisors ask, it builds dependence.'
)

doc.add_heading('Positive Observations', level=3)

doc.add_paragraph(
    'Most TC companies have Non-Conformance Report systems. Almost none have a system for recording '
    'positive observations. Build one. "Observed: TC repositioned taper approach due to wind gusts '
    'exceeding 40km/h. Adjusted cone spacing to next standard up. Maintained traffic flow. Good '
    'dynamic assessment." This does two things: it creates a record of good practice (evidence for '
    'the organisation), and it signals to the crew that someone notices when they do the right thing, '
    'not just when they do the wrong thing.'
)

doc.add_heading('The Close Call Log', level=3)

doc.add_paragraph(
    '"Near-miss reporting" carries negative connotation \u2014 something almost went wrong. Reframe '
    'it as a Close Call Log \u2014 something went right because someone noticed. "Close call: Wind '
    'started blowing cones into live lane. TC on east approach noticed and weighted cones before any '
    'entered traffic. No incident." The framing is critical: the good behaviour prevented the incident. '
    'The close call log documents the presence of safety capacity, not the absence of an incident.'
)

doc.add_heading('Crew Self-Assessment', level=3)

doc.add_paragraph(
    'After each job, the crew rates themselves on the 3+1. Not a supervisor rating \u2014 self-rating. '
    'Competent Crew: 4 out of 5. Sound Plan: 5 out of 5. Right Equipment: 3 out of 5 \u2014 weighted '
    'bases were at other depot. Dynamic Awareness: 4 out of 5 \u2014 caught the wind shift early but '
    'missed the sun glare until it was already a problem. This is self-monitoring, which is the '
    'Independent stage on the Bradley Curve. When crews start sharing their self-assessments with each '
    'other and learning across crews, that is Interdependent.'
)

doc.add_heading('Peer Observation', level=3)

doc.add_paragraph(
    'DuPont-style behavioural observation has a supervisor watching workers. Safety II flips this: TCs '
    'observe each other. Not to catch mistakes \u2014 to learn from each other\'s practice. "I watched '
    'how you managed that queue \u2014 I\'d have done it differently. Can I ask why you chose that '
    'position?" This normalises conversation about decision-making. It makes Dynamic Awareness '
    'discussable. The key difference: peer observation is horizontal (crew to crew), not vertical '
    '(management to worker). It builds interdependence. Management observation builds dependence.'
)

doc.add_heading('6.3 The Technology Opportunity', level=2)

doc.add_paragraph(
    'Current safety technology in the TC industry focuses almost entirely on compliance \u2014 is the '
    'TMP signed, is the SWMS on site, is the audit complete. No existing system captures evidence of '
    'Dynamic Awareness in real time. This represents a significant opportunity. The TC Work Zone '
    'Locator app, for example, already has capabilities that could be directed toward making good '
    'safety behaviour visible:'
)

headers = ['App Capability', 'What It Captures', 'Good Behaviour Evidence']
rows = [
    ['GPS + speed zone lookahead', 'TC checks speed zone before arrival', 'Proactive verification (Pillar 2)'],
    ['Live weather feed', 'Crew notes wind change, adjusts setup', 'Dynamic Awareness \u2014 reading conditions'],
    ['Banner alerts near SLK', 'Crew reviews nearby alerts during prestart', 'Dynamic Awareness \u2014 learning from others'],
    ['AfterCare lookahead', 'Crew positions with awareness of downstream', 'Dynamic Awareness \u2014 anticipating'],
    ['Close call logging (new)', 'TC logs a close call in real time', 'Evidence of active monitoring and response'],
]
add_styled_table(doc, headers, rows)

doc.add_paragraph('')
doc.add_paragraph(
    'Nobody in the TC industry is capturing this data. The first system that does will have a '
    'significant advantage \u2014 not just as a compliance tool, but as a safety culture tool that '
    'makes good practice visible and measurable for the first time.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 7. WHERE THIS LIVES IN EXISTING STANDARDS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('7. Where This Lives in Existing Standards', level=1)

doc.add_paragraph(
    'The framework does not introduce new obligations. Every element maps to existing requirements in '
    'Australian and Western Australian standards. What the framework provides is structure \u2014 a way '
    'to organise existing obligations so that the gap between planning and practice becomes visible.'
)

headers = ['Framework Element', 'Existing Standard Reference', 'What the Framework Adds']
rows = [
    ['Pillar 1:\nCompetent Crew',
     'AGTTM Part 7 \u2014 TC competency requirements\nMRWA CoP Section 8 \u2014 Accreditation\nWHS Act S27 \u2014 Officer due diligence',
     'Distinguishes accreditation from competence. Makes context-specific experience visible, not just card status.'],
    ['Pillar 2:\nSound Plan',
     'MRWA CoP Section 4 \u2014 TMP requirements\nAGTTM Part 3 \u2014 TGS design\nTMP 0922 Section 3.2.1 \u2014 Plan adequacy',
     'Requires the plan to be verified as current and site-specific, not just present. Adds crew understanding as a criterion.'],
    ['Pillar 3:\nRight Equipment',
     'MRWA CoP Section 6.16 \u2014 Speed feedback signs\nMRWA CoP Section 6.3 \u2014 PPE\nMRWA CoP Section 6.14 \u2014 Implementation and maintenance',
     'Requires serviceability and sufficiency, not just presence. Adds contingency equipment.'],
    ['Practice:\nDynamic Awareness',
     'TMP 0922 Section 3.2.1.2 \u2014 Site adjustment authority\nTMP 0922 Section 3.2.2 \u2014 Environmental conditions\nWHS Act S19 \u2014 Primary duty of care',
     'Names and structures what is currently one sentence in a risk register. Provides the three-component model (read, adjust, stop). Grounds it in Safety II.'],
    ['Thread:\nCommunication',
     'MRWA CoP Section 6.14 \u2014 TTM operation\nAGTTM Part 7 \u2014 Communication protocols\nWHS Act S19 \u2014 Duty to consult',
     'Positions communication as the enabler of all other elements, not a standalone compliance item.'],
]
add_styled_table(doc, headers, rows)

doc.add_paragraph('')
doc.add_paragraph(
    'The key mapping is between Dynamic Awareness and TMP Section 3.2.1.2, which states that Traffic '
    'Management Personnel may adjust or change the site where a significant safety hazard is identified. '
    'This single sentence is the regulatory authorisation for dynamic awareness. The framework '
    'transforms it from a buried provision into a visible, teachable, measurable practice.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 8. THE MATURITY MODEL
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('8. The Maturity Model', level=1)

doc.add_paragraph(
    'Not all organisations practise the framework equally. The maturity model describes four levels '
    'of adoption, drawing on the Bradley Curve used in industrial safety culture assessment. It '
    'provides organisations with a way to assess where they currently sit and what the next step looks like.'
)

headers = ['Level', 'Name', 'Pillars Present', 'Practice Present', 'What This Looks Like']
rows = [
    ['1', 'Compliance Only',
     '2 and 3\n(signs and plan)',
     'No',
     'The minimum. Signs are loaded, the TMP is on the dashboard. Crew competency is assumed based on accreditation. No process for reassessment during the job.'],
    ['2', 'Competent\nIndividuals',
     '1, 2, and 3\n(crew, plan, equipment)',
     'Informal only',
     'Qualified crew doing the right thing individually. Experienced TCs may practise dynamic awareness personally, but it\'s not shared or supported by the organisation.'],
    ['3', 'Coordinated\nTeam',
     '1, 2, and 3\n(all verified)',
     'Partially\n(no stop authority)',
     'Crew communicates and coordinates effectively. Conditions are discussed. But stop authority is unclear or conditional \u2014 TCs hesitate to call a halt.'],
    ['4', 'Adaptive\nCrew',
     '1, 2, and 3\n(all verified)',
     'Yes\n(read, adjust, stop)',
     'The full framework. The team reads conditions, adjusts controls when needed, and has clear stop authority. This is not dependent on individual experience \u2014 it is embedded in the crew\'s culture.'],
]
add_styled_table(doc, headers, rows)

doc.add_paragraph('')
doc.add_paragraph(
    'An organisation at Level 2 can pass an audit. An organisation at Level 4 is where nobody gets hurt. '
    'The difference is not in the paperwork \u2014 it is in the practice.'
)

doc.add_paragraph(
    'Most traffic management companies in Western Australia currently operate between Level 1 and '
    'Level 2. They have the signs and the plan. Some have competent crews. Very few have systematically '
    'addressed dynamic awareness as a safety practice, and even fewer have explicitly granted stop '
    'authority to their Traffic Controllers.'
)

doc.add_paragraph(
    'The maturity model also maps to the Bradley Curve stages. Level 1 corresponds to the Dependent '
    'stage \u2014 people follow rules because they are told to. Level 2 corresponds to early Independent '
    '\u2014 individuals take personal responsibility but don\'t share it. Level 3 is late Independent '
    '\u2014 the team coordinates but lacks the cultural permission to challenge. Level 4 is '
    'Interdependent \u2014 the team monitors itself, adapts collectively, and has full stop authority. '
    'Dynamic Awareness is the behavioural expression of the Interdependent stage. It cannot be achieved '
    'through compliance alone; it requires a cultural shift that the organisation must lead.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 9. WHAT THIS MEANS FOR ORGANISATIONS
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('9. What This Means for Organisations', level=1)

doc.add_heading('9.1 The Two Layers', level=2)

doc.add_paragraph(
    'The framework operates at two layers. The crew layer is what the Traffic Controller can control '
    'and practice: verifying the three pillars at the pre-start, practising dynamic awareness throughout '
    'the shift, and communicating continuously. The organisational layer is what the company must provide: '
    'the systems, resourcing, and culture that make the crew layer possible.'
)

doc.add_paragraph(
    'The banner alerts demonstrate what happens when the organisational layer fails. In EQ#56820, the '
    'prestart form was reused for a week because the company\'s system allowed it. In EQ#57801, proper '
    'recovery equipment was not available because the company did not equip the crew for contingencies. '
    'In each case, the crew could not practise what the framework requires because the organisation had '
    'not provided the conditions for it.'
)

headers = ['Layer', 'Responsibility', 'Examples']
rows = [
    ['Crew', 'What the TC can control and practice',
     'Verifying pillars at pre-start\nReading conditions during the job\nAdjusting controls when conditions change\nCalling a stop when something feels wrong\nCommunicating observations to the team'],
    ['Organisational', 'What the company must provide',
     'Systems that ensure plans are current and task-specific\nResourcing that puts competent supervision on every job\nCulture that grants stop authority without consequence\nLearning loops that feed incidents back into practice\nEquipment that goes beyond the minimum schedule\nCrew leaders who model Dynamic Awareness, not just compliance'],
]
add_styled_table(doc, headers, rows)

doc.add_heading('9.2 Breaking the Chain: Practical Interventions', level=2)

doc.add_paragraph(
    'The only way to stop cultural reproduction of bad practice is to interrupt the apprenticeship. '
    'The following interventions are designed to be practical, low-cost, and implementable without '
    'new systems or new paperwork:'
)

doc.add_heading('Separate "Buddy" from "Model"', level=3)

doc.add_paragraph(
    'Currently, new TCs are assigned to whichever crew has a spare spot. The person they learn from '
    'is random. Instead, designate mentors \u2014 experienced TCs specifically selected and trained to '
    'model Dynamic Awareness. Mentor criteria should not be "has been doing this for five years" but '
    '"can describe their decision-making process in real time" and "has demonstrated Dynamic Awareness '
    'under observation." New TCs should spend their first 20 jobs with mentors only. This breaks the '
    'transmission chain because the first 20 jobs establish the normal.'
)

doc.add_heading('Make the Invisible Visible', level=3)

doc.add_paragraph(
    'The new TC copies the crew leader because they assume the crew leader is the standard. Make the '
    'actual standard visible through briefing cards \u2014 simple, pocket-sized cards that describe '
    'what Dynamic Awareness looks like in practice: "Scan the approach every 30 seconds. Check the sky '
    'for weather changes. Listen for speed changes. Reposition if conditions shift." These are not '
    'checklists \u2014 they are descriptions of behaviour, not lists of tasks. The difference matters. '
    'A checklist gets ticked. A behaviour description gets practised.'
)

doc.add_heading('Raise the Crew Leader Standard', level=3)

doc.add_paragraph(
    'Not everyone with a traffic management ticket should be a crew leader. Currently, crew leader '
    'selection is based on having the right tickets, being available, and not being the newest person. '
    'It should be based on whether they can model Dynamic Awareness, can explain their decision-making '
    'to a new TC, can run a prestart that is a conversation rather than a monologue, and can handle a '
    'situation that is not in the TMP. If the answer to "will new TCs be better after working with you '
    'or will they learn your shortcuts?" is the latter, that person should not be crew leading.'
)

doc.add_heading('The "What Did You See?" Protocol', level=3)

doc.add_paragraph(
    'Every crew, every job, every handover: the outgoing crew leader asks the incoming crew leader '
    '"What did you see today?" Not "how was it?" \u2014 that gets "fine." "What did you see?" requires '
    'a specific answer about observations, adjustments, and changes. The new TC hears this exchange '
    'and learns that talking about what you noticed is normal. It takes 30 seconds. It is not a '
    'meeting. It is not paperwork. It is a conversation that models the behaviour you want.'
)

doc.add_heading('9.3 Implementation Without Burden', level=2)

doc.add_paragraph(
    'The framework is designed to be adopted without adding compliance burden. It does not require new '
    'forms, new checklists, or new reporting obligations. It requires a shift in how existing processes '
    'are used:'
)

add_bullet(doc, 'The pre-start brief already happens. The framework asks that it covers all three pillars explicitly, not just the plan.', 'Pre-start briefs: ')
add_bullet(doc, 'Site adjustments are already permitted under TMP 3.2.1.2. The framework asks that they are made proactively, not only after an incident.', 'Site adjustments: ')
add_bullet(doc, 'Stop authority is already implicit in the WHS duty of care. The framework asks that it is made explicit \u2014 every crew member knows they can call a stop, and every supervisor backs them up.', 'Stop authority: ')
add_bullet(doc, 'Banner alerts are already issued by MRWA. The framework asks that they are used as learning tools, not just filed as compliance documents.', 'Incident learning: ')

doc.add_paragraph(
    'The shift is cultural, not procedural. The three pillars are already being checked in most '
    'organisations, though often implicitly. The practice of dynamic awareness is already being performed '
    'by experienced TCs, though often unconsciously. The framework makes both visible, teachable, and '
    'measurable \u2014 without requiring a new system to do it.'
)

doc.add_heading('9.4 The Business Case', level=2)

doc.add_paragraph(
    'For organisations, the framework offers a clear business case beyond the moral imperative of '
    'keeping workers safe:'
)

add_bullet(doc, 'Every Lost Time Injury costs the organisation in workers\' compensation, lost productivity, investigation time, and replacement labour. The banner alerts document LTIs that occurred on jobs where the preconditions were met \u2014 meaning these were preventable with dynamic awareness, not with more paperwork.', 'Reduced incidents: ')
add_bullet(doc, 'Organisations that can demonstrate a mature safety culture \u2014 including the practice of dynamic awareness and explicit stop authority \u2014 are better positioned for MRWA pre-qualification, contract evaluation, and insurance assessment.', 'Regulatory positioning: ')
add_bullet(doc, 'TCs who feel empowered to speak up and adjust controls are more engaged, more likely to stay with the organisation, and more likely to develop into supervisors. The framework supports retention by valuing the judgement that experienced TCs already exercise.', 'Crew retention: ')
add_bullet(doc, 'The WHS Act 2020 Section 27 requires officers to exercise due diligence, including acquiring and keeping up-to-date knowledge of hazards and risks. Dynamic awareness is the operational expression of that duty \u2014 it is how an organisation demonstrates that it is managing risks in real time, not just on paper.', 'Due diligence: ')

# ═══════════════════════════════════════════════════════════════════════════════
# 10. THE EVIDENCE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('10. The Evidence: Banner Alerts Mapped to the Framework', level=1)

doc.add_paragraph(
    'The evidence base for this framework has been expanded from 7 to 64 MRWA Banner Alerts. '
    'Each alert has been classified as either directly or indirectly related to traffic control (TC) '
    'for filtering purposes. This classification allows the framework to distinguish between incidents '
    'that occur within the traffic control system itself and those that occur in the broader work '
    'environment but still affect TC workers.'
)

add_horizontal_rule(doc)

doc.add_heading('10.1 Overview of the Evidence Base', level=2)

doc.add_paragraph(
    'The 64 banner alerts span the period from November 2023 to April 2026 and include the full '
    'range of incident types reported through the MRWA Banner Alert system. The following key '
    'statistics describe the dataset:'
)

add_bullet(doc, '64 (45 Red, 19 Grey, 0 Amber)', 'Total alerts: ')
add_bullet(doc, '25 (39.1%)', 'Directly related to TC: ')
add_bullet(doc, '39 (60.9%)', 'Indirectly related to TC: ')
add_bullet(doc, '3', 'LTI count in direct alerts: ')

add_callout(doc,
    'The direct/indirect classification is not a judgement on severity. Indirect alerts can and do '
    'result in serious injuries. The classification is a filtering tool designed to identify which '
    'alerts speak most directly to the traffic control system itself, and therefore which alerts are '
    'most relevant to validating the 3 Pillars, 1 Practice framework.',
    label='Note'
)

doc.add_heading('10.2 Direct TC Alert Subcategories', level=2)

doc.add_paragraph(
    'The 25 directly TC-related alerts fall into five subcategories, each representing a distinct '
    'type of interaction with the traffic control system:'
)

headers = ['Subcategory', 'Count', 'Description']
rows = [
    ['mop_breach', '10', 'Member of Public breaching or interacting with TC setup/work zone'],
    ['tm_breach', '8', 'Traffic management breach (1.2m rule, wrong TGS, TM procedure violation)'],
    ['tc_injured', '4', 'Traffic Controller injured while performing TC duties'],
    ['tc_equipment', '2', 'TC equipment involved (TMA, follow-me vehicle)'],
    ['tm_setup', '1', 'TM setup/removal activity (placing signage, entering live lane)'],
]
add_styled_table(doc, headers, rows)

doc.add_paragraph('')

doc.add_heading('10.3 Indirect TC Alert Subcategories', level=2)

doc.add_paragraph(
    'The 39 indirectly TC-related alerts fall into eleven subcategories, representing hazards that '
    'affect TC workers but are not inherent to the traffic control system itself:'
)

headers = ['Subcategory', 'Count', 'Description']
rows = [
    ['journey', '7', 'Journey management incident (driving to/from site)'],
    ['equipment_tool', '6', 'Equipment or tool injury'],
    ['slip_trip', '5', 'Slips, trips, falls on site'],
    ['manual_handling', '5', 'Manual handling injury (back strain, lifting, pulling)'],
    ['fitness_for_duty', '4', 'Fitness for duty (drug/alcohol, medical episode, fatigue)'],
    ['mechanical', '3', 'Mechanical failure (leaf spring, chain, vehicle component)'],
    ['road_furniture', '3', 'Road furniture maintenance (guide posts, wire ropes)'],
    ['construction', '2', 'Construction activity injury'],
    ['vehicle_incident', '2', 'Vehicle incident not in work zone or TM-related'],
    ['utility_strike', '1', 'Utility strike (power line, service)'],
    ['environmental', '1', 'Environmental hazard (thermal stress, weather)'],
]
add_styled_table(doc, headers, rows)

doc.add_paragraph('')

doc.add_heading('10.4 Detailed Direct TC Alert Register', level=2)

doc.add_paragraph(
    'The following table lists all 25 directly TC-related banner alerts, including their '
    'classification, subcategory, description, and LTI status:'
)

headers = ['EQ#', 'Date', 'Banner', 'Subcategory', 'Description', 'LTI']
rows = [
    ['EQ#35211', '2023-11-08', 'Grey', 'tm_breach', 'TC on foot within 1.2m in a 60km speed zone', 'No'],
    ['EQ#35216', '2023-11-13', 'Grey', 'mop_breach', 'MoP collided with MR Employee after veering into oncoming lane', 'No'],
    ['EQ#37100', '2024-01-20', 'Red', 'tm_breach', 'Mobile Plant and Traffic Management Procedure Breach', 'No'],
    ['EQ#37292', '2024-01-30', 'Red', 'tc_injured', 'TC found lying face down with concussion', 'Yes'],
    ['EQ#37443', '2024-02-04', 'Grey', 'tc_equipment', 'Follow me vehicle conducts U-turn in front of road train', 'No'],
    ['EQ#37689', '2024-02-08', 'Red', 'mop_breach', 'MoP driven through Vehicle Control Point colliding with TC and Fire Services', 'No'],
    ['EQ#37835', '2024-02-24', 'Red', 'tc_equipment', 'Lower back strain accessing TMA resulting in LTI', 'Yes'],
    ['EQ#38338', '2024-02-23', 'Red', 'mop_breach', 'MoP Heavy vehicle loses load and damages TC light vehicle', 'No'],
    ['EQ#38369', '2024-03-22', 'Red', 'mop_breach', 'Worker struck by Member of Public', 'No'],
    ['EQ#38716', '2024-03-08', 'Red', 'tm_breach', 'Breach of Mobile Plan Onboarding Procedure', 'No'],
    ['EQ#38753', '2024-03-11', 'Red', 'mop_breach', 'Member of Public Rollover within Works Area', 'No'],
    ['EQ#39231', '2024-03-28', 'Grey', 'tm_breach', 'TM Breach - Lab (core) Tester working within 1.2m of live traffic', 'No'],
    ['EQ#39234', '2024-03-28', 'Red', 'tm_breach', 'TM Breach - Lab (core) Tester working within 1.2m of live traffic', 'No'],
    ['EQ#39486', '2024-04-09', 'Grey', 'tm_breach', 'TC and work vehicle within 1.2m of live traffic lane', 'No'],
    ['EQ#40545', '2024-05-21', 'Red', 'tm_breach', 'Core tester observed working within 1.2m of live traffic at 110km/hr', 'No'],
    ['EQ#40615', '2024-05-24', 'Red', 'mop_breach', 'MoP collides with Contractor Vehicle turning into Site Office', 'No'],
    ['EQ#40656', '2024-05-30', 'Red', 'tc_injured', 'TC experiences back pain whilst performing manual labour tasks', 'No'],
    ['EQ#41100', '2024-06-11', 'Red', 'mop_breach', 'OSOM breaches traffic management', 'No'],
    ['EQ#41942', '2024-06-26', 'Red', 'tc_injured', 'TC working excessive hours without sufficient breaks between working shifts', 'No'],
    ['EQ#42098', '2024-07-02', 'Red', 'mop_breach', 'Learner Driver overtaking causes maintenance truck to veer from road', 'No'],
    ['EQ#43968', '2024-09-05', 'Grey', 'tm_breach', 'Incorrect Traffic Guidance Scheme', 'No'],
    ['EQ#47857', '2025-02-02', 'Red', 'mop_breach', 'Heavy Vehicle veers into work area', 'No'],
    ['EQ#48009', '2025-02-04', 'Red', 'mop_breach', 'Near miss from light vehicle overtaking 3 stationary trucks and running a Red Light', 'No'],
    ['EQ#48096', '2025-02-12', 'Red', 'tm_setup', 'Worker enters live to place traffic signage', 'No'],
    ['EQ#57872', '2026-04-11', 'Red', 'tc_injured', 'TC sustains shoulder injury slipping on roadside batter', 'Yes'],
]
add_styled_table(doc, headers, rows)

doc.add_paragraph('')

doc.add_heading('10.5 Analysis: What the Evidence Tells Us', level=2)

doc.add_heading('MoP Breaches (10 incidents)', level=3)

doc.add_paragraph(
    'The single largest category of TC-direct incidents involves Members of Public interacting '
    'with or breaching traffic control setups. This is significant because it is the category of '
    'risk that compliance systems are least equipped to manage. A TMP cannot control driver '
    'behaviour. A TGS cannot prevent a motorist from driving through a Vehicle Control Point '
    '(EQ#37689) or veering into oncoming traffic (EQ#35216). These incidents validate the '
    'framework\'s emphasis on Dynamic Awareness: the crew must continuously read the traffic '
    'behaviour around them and adjust controls in real time. The plan assumes compliance from '
    'the public; Dynamic Awareness prepares for when that assumption fails.'
)

doc.add_heading('TM Breaches (8 incidents)', level=3)

doc.add_paragraph(
    'The second largest category involves breaches of traffic management procedures \u2014 workers '
    'within 1.2m of live traffic, incorrect TGS implementations, and onboarding procedure failures. '
    'Critically, several of these breaches occurred not because the workers were reckless, but because '
    'of systemic pressures: time pressure to complete core testing before the rostered swing ended '
    '(EQ#39231, EQ#39234), inadequate supervision when project engineers departed site early '
    '(EQ#39231), and crew not fully aware of hazards or how to implement effective controls as '
    'specified in the SWMS (EQ#39486). These are not individual failures \u2014 they are system '
    'failures that individual compliance cannot prevent. The 1.2m rule appears in 5 of the 8 TM '
    'breaches, suggesting this is a critical risk area where the gap between the rule and the reality '
    'of getting work done near live traffic is widest.'
)

doc.add_heading('TC Injuries (4 incidents)', level=3)

doc.add_paragraph(
    'Four incidents involve Traffic Controllers injured while performing their duties \u2014 a TC '
    'found lying face down with concussion (EQ#37292), a TC experiencing back pain performing manual '
    'labour (EQ#40656), a TC working excessive hours without sufficient breaks (EQ#41942), and a TC '
    'sustaining a shoulder injury slipping on a roadside batter (EQ#57872). Three of these four '
    'resulted in LTI. The pattern here is telling: TCs are being injured by the physical demands '
    'and environmental conditions of the role, not by the traffic they are controlling. The TC who '
    'was found face down with concussion \u2014 the alert does not specify how the injury occurred, '
    'which itself is an indicator of the isolated nature of TC work. Dynamic Awareness, in this '
    'context, means crew members watching out for each other, monitoring fatigue, and recognising '
    'when a colleague is at risk.'
)

doc.add_heading('TC Equipment and TM Setup (3 incidents)', level=3)

doc.add_paragraph(
    'Two incidents involve TC-specific equipment \u2014 a follow-me vehicle conducting a U-turn in '
    'front of a road train (EQ#37443) and a TC sustaining a back injury accessing a Truck Mounted '
    'Attenuator (EQ#37835). One incident involves a worker entering a live lane to place traffic '
    'signage (EQ#48096). These incidents highlight the risks inherent in the setup and pack-down '
    'phases of traffic management, which are often the most hazardous moments of a shift because '
    'the full TM layout is not yet in place or has been partially removed.'
)

doc.add_heading('The Indirect Evidence', level=3)

doc.add_paragraph(
    'The 39 indirectly related alerts provide important context. Journey management incidents (7) '
    'represent the largest indirect category, highlighting that the risks to TC workers extend beyond '
    'the work zone itself \u2014 the drive to and from remote sites is a significant hazard. Fitness '
    'for duty incidents (4), including positive BAC results and a driver who fell asleep at the wheel, '
    'underscore the importance of Pillar 1 (Competent Crew) being about more than just accreditation '
    '\u2014 it includes being fit for work. Manual handling and slip/trip incidents (10 combined) '
    'point to the physical nature of TC work and the need for environmental awareness that goes '
    'beyond traffic management.'
)

doc.add_heading('Summary', level=3)

doc.add_paragraph(
    'The evidence from 64 banner alerts confirms what the framework predicts: the three pillars '
    '(verified at pre-start) are necessary but not sufficient. The 25 directly TC-related alerts '
    'show that the most common failure modes \u2014 MoP breaches, TM procedure breaches, TC injuries, '
    'and equipment/setup incidents \u2014 all involve the gap between what the plan assumes and what '
    'happens in reality. Dynamic Awareness, the continuous practice of reading conditions and '
    'adjusting controls, is the framework\'s response to this evidence.'
)

# ═══════════════════════════════════════════════════════════════════════════════
# 11. THE PATH FORWARD
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_heading('11. The Path Forward', level=1)

doc.add_paragraph(
    'This paper is presented for industry consultation. The framework is offered as a structure for '
    'a conversation that many in traffic management are already having informally: how do we keep '
    'people safe when the plan isn\'t enough? The 3 Pillars, 1 Practice framework does not claim to '
    'have all the answers. It claims to have named the right question.'
)

doc.add_paragraph(
    'The framework sits at the intersection of two safety traditions. The three pillars draw on the '
    'proven Safety I approach of verification, compliance, and control. The practice of Dynamic '
    'Awareness draws on the emerging Safety II paradigm of adaptive capacity, learning from success, '
    'and treating workers as the resource rather than the problem. Neither tradition alone is sufficient '
    'for the unique conditions of temporary traffic management. Together, they address the structural '
    'contradiction at the heart of the industry: we apply Safe System thinking to road users but '
    'compliance-based thinking to the workers who implement it.'
)

doc.add_paragraph(
    'The practical interventions proposed in this paper \u2014 designated mentors, the "What Did You '
    'See?" protocol, positive observation systems, crew self-assessment, and close call logging \u2014 '
    'are designed to be implementable without new systems, new paperwork, or new compliance burden. '
    'They require a cultural shift: from monitoring compliance to making good practice visible, from '
    'auditing what went wrong to studying what goes right, and from holding individuals accountable '
    'for minimums to holding organisations accountable for creating the conditions where the minimum '
    'is not the ceiling.'
)

doc.add_paragraph(
    'Feedback is welcomed from Traffic Controllers, supervisors, safety managers, and regulatory '
    'bodies. The framework will be refined through industry input before any formal proposal is '
    'submitted to Main Roads Western Australia.'
)

# ─── Save ─────────────────────────────────────────────────────────────────────

output_path = '/home/z/my-project/download/Beyond_the_Plan_3_Pillars_1_Practice_Framework.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
